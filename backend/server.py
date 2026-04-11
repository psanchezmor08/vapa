from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
from datetime import datetime, timezone, timedelta
from pathlib import Path
import os
import logging
import uuid
import jwt
from passlib.context import CryptContext
import ipaddress
import qrcode
import io
import base64
import secrets
import string
import hashlib
import json
from urllib.parse import quote, unquote

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = os.environ.get('SECRET_KEY', 'vapa-secret-key-change-in-production')
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24  # 24 hours

security = HTTPBearer()

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# ============= MODELS =============

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    password_hash: str
    role: str = "viewer"  # admin, editor, viewer
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    role: str = "viewer"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    role: str
    created_at: datetime

class Token(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse

class BlogPost(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    slug: str
    content: str
    excerpt: str
    author_id: str
    published: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class BlogPostCreate(BaseModel):
    title: str
    content: str
    excerpt: str
    published: bool = True

class BlogPostUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    excerpt: Optional[str] = None
    published: Optional[bool] = None

class SubnetCalculatorInput(BaseModel):
    ip_address: str
    cidr: int

class SubnetCalculatorOutput(BaseModel):
    network_address: str
    broadcast_address: str
    first_usable_ip: str
    last_usable_ip: str
    total_hosts: int
    usable_hosts: int
    subnet_mask: str
    wildcard_mask: str
    cidr: int
    ip_class: str
    network_binary: str
    subnet_binary: str

class QRGeneratorInput(BaseModel):
    text: str
    size: int = 300

class QRGeneratorOutput(BaseModel):
    qr_code: str  # Base64 encoded image

class PasswordGeneratorInput(BaseModel):
    length: int = 16
    use_uppercase: bool = True
    use_lowercase: bool = True
    use_digits: bool = True
    use_symbols: bool = True

class PasswordGeneratorOutput(BaseModel):
    password: str
    strength: str

class Base64Input(BaseModel):
    text: str
    encode: bool = True  # True for encode, False for decode

class Base64Output(BaseModel):
    result: str

class HashInput(BaseModel):
    text: str
    algorithm: str = "sha256"  # md5, sha1, sha256, sha512

class HashOutput(BaseModel):
    hash: str
    algorithm: str

class JSONValidatorInput(BaseModel):
    json_string: str

class JSONValidatorOutput(BaseModel):
    valid: bool
    formatted: Optional[str] = None
    error: Optional[str] = None

class UnitConverterInput(BaseModel):
    value: float
    from_unit: str  # bytes, kb, mb, gb, tb
    to_unit: str

class UnitConverterOutput(BaseModel):
    result: float
    from_unit: str
    to_unit: str

class UUIDGeneratorOutput(BaseModel):
    uuid: str

class PasswordChange(BaseModel):
    current_password: str
    new_password: str

class PasswordReset(BaseModel):
    new_password: str

class UserUpdate(BaseModel):
    role: Optional[str] = None

class URLEncoderInput(BaseModel):
    text: str
    encode: bool = True

class URLEncoderOutput(BaseModel):
    result: str

class PortAnalyzerInput(BaseModel):
    port: int

class PortAnalyzerOutput(BaseModel):
    port: int
    service: str
    description: str

# ============= AUTH HELPERS =============

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.now(timezone.utc) + expires_delta
    else:
        expire = datetime.now(timezone.utc) + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> User:
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid authentication credentials")
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.JWTError:
        raise HTTPException(status_code=401, detail="Could not validate credentials")
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    return User(**user)

async def get_current_active_editor(current_user: User = Depends(get_current_user)) -> User:
    if current_user.role not in ["admin", "editor"]:
        raise HTTPException(status_code=403, detail="No tienes permisos suficientes")
    return current_user

async def get_current_active_admin(current_user: User = Depends(get_current_user)) -> User:
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Se requiere rol de administrador")
    return current_user
# ============= AUTH ROUTES =============

@api_router.post("/auth/register", response_model=UserResponse)
async def register(user_data: UserCreate, current_user: User = Depends(get_current_active_admin)):
    # Check if user exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    user = User(
        email=user_data.email,
        password_hash=get_password_hash(user_data.password),
        role=user_data.role
    )
    
    doc = user.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.users.insert_one(doc)
    
    return UserResponse(**user.model_dump())

@api_router.post("/auth/login", response_model=Token)
async def login(user_data: UserLogin):
    user = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if not user or not verify_password(user_data.password, user['password_hash']):
        raise HTTPException(status_code=401, detail="Incorrect email or password")
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user['id']}, expires_delta=access_token_expires
    )
    
    if isinstance(user['created_at'], str):
        user['created_at'] = datetime.fromisoformat(user['created_at'])
    
    return Token(
        access_token=access_token,
        token_type="bearer",
        user=UserResponse(**user)
    )
@api_router.get("/users", response_model=List[UserResponse])
async def get_users(current_user: User = Depends(get_current_active_admin)):
    users = await db.users.find({}, {"_id": 0}).to_list(100)
    result = []
    for u in users:
        if isinstance(u['created_at'], str):
            u['created_at'] = datetime.fromisoformat(u['created_at'])
        result.append(UserResponse(**u))
    return result

@api_router.post("/users", response_model=UserResponse)
async def create_user(user_data: UserCreate, current_user: User = Depends(get_current_active_admin)):
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    user = User(
        email=user_data.email,
        password_hash=get_password_hash(user_data.password),
        role=user_data.role
    )
    doc = user.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.users.insert_one(doc)
    return UserResponse(**user.model_dump())

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(get_current_active_admin)):
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="No puedes eliminarte a ti mismo")
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")
    return {"message": "Usuario eliminado"}

@api_router.put("/users/{user_id}/password")
async def reset_user_password(user_id: str, data: PasswordReset, current_user: User = Depends(get_current_active_admin)):
    result = await db.users.update_one(
        {"id": user_id},
        {"$set": {"password_hash": get_password_hash(data.new_password)}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")
    return {"message": "Contraseña restablecida"}

@api_router.put("/auth/change-password")
async def change_my_password(data: PasswordChange, current_user: User = Depends(get_current_user)):
    if not verify_password(data.current_password, current_user.password_hash):
        raise HTTPException(status_code=400, detail="Contraseña actual incorrecta")
    await db.users.update_one(
        {"id": current_user.id},
        {"$set": {"password_hash": get_password_hash(data.new_password)}}
    )
    return {"message": "Contraseña cambiada correctamente"}

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: User = Depends(get_current_user)):
    return UserResponse(**current_user.model_dump())

# ============= BLOG ROUTES =============

@api_router.get("/blog/posts", response_model=List[BlogPost])
async def get_blog_posts(published_only: bool = True, credentials: Optional[HTTPAuthorizationCredentials] = Depends(HTTPBearer(auto_error=False))):
    user_role = None
    if credentials:
        try:
            payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
            user_id = payload.get("sub")
            if user_id:
                u = await db.users.find_one({"id": user_id}, {"_id": 0})
                if u:
                    user_role = u.get("role")
        except:
            pass
    if user_role in ["admin", "editor", "viewer"]:
        query = {} if not published_only else {"published": True}
    else:
        query = {"published": True}
    posts = await db.blog_posts.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    for post in posts:
        if isinstance(post['created_at'], str):
            post['created_at'] = datetime.fromisoformat(post['created_at'])
        if isinstance(post['updated_at'], str):
            post['updated_at'] = datetime.fromisoformat(post['updated_at'])
    return posts

@api_router.get("/blog/posts/{post_id}", response_model=BlogPost)
async def get_blog_post(post_id: str):
    post = await db.blog_posts.find_one({"id": post_id}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    
    if isinstance(post['created_at'], str):
        post['created_at'] = datetime.fromisoformat(post['created_at'])
    if isinstance(post['updated_at'], str):
        post['updated_at'] = datetime.fromisoformat(post['updated_at'])
    
    return BlogPost(**post)

@api_router.post("/blog/posts", response_model=BlogPost)
async def create_blog_post(post_data: BlogPostCreate, current_user: User = Depends(get_current_active_editor)):
    # Generate slug from title
    slug = post_data.title.lower().replace(" ", "-").replace(":", "").replace("?", "")
    
    post = BlogPost(
        title=post_data.title,
        slug=slug,
        content=post_data.content,
        excerpt=post_data.excerpt,
        author_id=current_user.id,
        published=post_data.published
    )
    
    doc = post.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['updated_at'] = doc['updated_at'].isoformat()
    await db.blog_posts.insert_one(doc)
    
    return post

@api_router.put("/blog/posts/{post_id}", response_model=BlogPost)
async def update_blog_post(post_id: str, post_data: BlogPostUpdate, current_user: User = Depends(get_current_active_editor)):
    post = await db.blog_posts.find_one({"id": post_id}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    
    update_data = post_data.model_dump(exclude_unset=True)
    if update_data:
        if 'title' in update_data:
            update_data['slug'] = update_data['title'].lower().replace(" ", "-").replace(":", "").replace("?", "")
        update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
        await db.blog_posts.update_one({"id": post_id}, {"$set": update_data})
    
    updated_post = await db.blog_posts.find_one({"id": post_id}, {"_id": 0})
    if isinstance(updated_post['created_at'], str):
        updated_post['created_at'] = datetime.fromisoformat(updated_post['created_at'])
    if isinstance(updated_post['updated_at'], str):
        updated_post['updated_at'] = datetime.fromisoformat(updated_post['updated_at'])
    
    return BlogPost(**updated_post)

@api_router.delete("/blog/posts/{post_id}")
async def delete_blog_post(post_id: str, current_user: User = Depends(get_current_active_editor)):
    result = await db.blog_posts.delete_one({"id": post_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Post not found")
    return {"message": "Post deleted successfully"}

# ============= TOOLS ROUTES =============

@api_router.post("/tools/subnet-calculator", response_model=SubnetCalculatorOutput)
async def calculate_subnet(input_data: SubnetCalculatorInput):
    try:
        network = ipaddress.ip_network(f"{input_data.ip_address}/{input_data.cidr}", strict=False)
        
        # Get IP class
        first_octet = int(input_data.ip_address.split('.')[0])
        if first_octet < 128:
            ip_class = "A"
        elif first_octet < 192:
            ip_class = "B"
        elif first_octet < 224:
            ip_class = "C"
        elif first_octet < 240:
            ip_class = "D (Multicast)"
        else:
            ip_class = "E (Reserved)"
        
        # Calculate addresses
        all_hosts = list(network.hosts())
        first_usable = str(all_hosts[0]) if len(all_hosts) > 0 else str(network.network_address)
        last_usable = str(all_hosts[-1]) if len(all_hosts) > 0 else str(network.broadcast_address)
        
        # Binary representations
        network_binary = '.'.join([bin(int(x))[2:].zfill(8) for x in str(network.network_address).split('.')])
        subnet_binary = '.'.join([bin(int(x))[2:].zfill(8) for x in str(network.netmask).split('.')])
        
        # Wildcard mask
        wildcard = ipaddress.IPv4Address(int(network.hostmask))
        
        return SubnetCalculatorOutput(
            network_address=str(network.network_address),
            broadcast_address=str(network.broadcast_address),
            first_usable_ip=first_usable,
            last_usable_ip=last_usable,
            total_hosts=network.num_addresses,
            usable_hosts=network.num_addresses - 2 if network.num_addresses > 2 else 0,
            subnet_mask=str(network.netmask),
            wildcard_mask=str(wildcard),
            cidr=input_data.cidr,
            ip_class=ip_class,
            network_binary=network_binary,
            subnet_binary=subnet_binary
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid IP or CIDR: {str(e)}")

@api_router.post("/tools/qr-generator", response_model=QRGeneratorOutput)
async def generate_qr(input_data: QRGeneratorInput):
    try:
        qr = qrcode.QRCode(version=1, box_size=10, border=5)
        qr.add_data(input_data.text)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to base64
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_str = base64.b64encode(buffer.getvalue()).decode()
        
        return QRGeneratorOutput(qr_code=f"data:image/png;base64,{img_str}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error generating QR: {str(e)}")

@api_router.post("/tools/password-generator", response_model=PasswordGeneratorOutput)
async def generate_password(input_data: PasswordGeneratorInput):
    try:
        characters = ""
        if input_data.use_lowercase:
            characters += string.ascii_lowercase
        if input_data.use_uppercase:
            characters += string.ascii_uppercase
        if input_data.use_digits:
            characters += string.digits
        if input_data.use_symbols:
            characters += string.punctuation
        
        if not characters:
            raise HTTPException(status_code=400, detail="At least one character type must be selected")
        
        password = ''.join(secrets.choice(characters) for _ in range(input_data.length))
        
        # Calculate strength
        strength = "Weak"
        if input_data.length >= 12 and sum([input_data.use_uppercase, input_data.use_lowercase, input_data.use_digits, input_data.use_symbols]) >= 3:
            strength = "Strong"
        elif input_data.length >= 8:
            strength = "Medium"
        
        return PasswordGeneratorOutput(password=password, strength=strength)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error generating password: {str(e)}")

@api_router.post("/tools/base64", response_model=Base64Output)
async def base64_converter(input_data: Base64Input):
    try:
        if input_data.encode:
            result = base64.b64encode(input_data.text.encode()).decode()
        else:
            result = base64.b64decode(input_data.text).decode()
        return Base64Output(result=result)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/hash", response_model=HashOutput)
async def hash_generator(input_data: HashInput):
    try:
        if input_data.algorithm == "md5":
            hash_obj = hashlib.md5(input_data.text.encode())
        elif input_data.algorithm == "sha1":
            hash_obj = hashlib.sha1(input_data.text.encode())
        elif input_data.algorithm == "sha256":
            hash_obj = hashlib.sha256(input_data.text.encode())
        elif input_data.algorithm == "sha512":
            hash_obj = hashlib.sha512(input_data.text.encode())
        else:
            raise HTTPException(status_code=400, detail="Invalid algorithm")
        
        return HashOutput(hash=hash_obj.hexdigest(), algorithm=input_data.algorithm)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/json-validator", response_model=JSONValidatorOutput)
async def json_validator(input_data: JSONValidatorInput):
    try:
        parsed = json.loads(input_data.json_string)
        formatted = json.dumps(parsed, indent=2, ensure_ascii=False)
        return JSONValidatorOutput(valid=True, formatted=formatted)
    except json.JSONDecodeError as e:
        return JSONValidatorOutput(valid=False, error=str(e))

@api_router.post("/tools/unit-converter", response_model=UnitConverterOutput)
async def unit_converter(input_data: UnitConverterInput):
    try:
        units = {"bytes": 1, "kb": 1024, "mb": 1024**2, "gb": 1024**3, "tb": 1024**4}
        
        from_bytes = input_data.value * units[input_data.from_unit.lower()]
        result = from_bytes / units[input_data.to_unit.lower()]
        
        return UnitConverterOutput(
            result=round(result, 4),
            from_unit=input_data.from_unit,
            to_unit=input_data.to_unit
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.get("/tools/uuid-generator", response_model=UUIDGeneratorOutput)
async def uuid_generator():
    return UUIDGeneratorOutput(uuid=str(uuid.uuid4()))

@api_router.post("/tools/url-encoder", response_model=URLEncoderOutput)
async def url_encoder(input_data: URLEncoderInput):
    try:
        if input_data.encode:
            result = quote(input_data.text)
        else:
            result = unquote(input_data.text)
        return URLEncoderOutput(result=result)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/port-analyzer", response_model=PortAnalyzerOutput)
async def port_analyzer(input_data: PortAnalyzerInput):
    common_ports = {
        20: ("FTP Data", "File Transfer Protocol (Data)"),
        21: ("FTP Control", "File Transfer Protocol (Control)"),
        22: ("SSH", "Secure Shell"),
        23: ("Telnet", "Telnet Protocol"),
        25: ("SMTP", "Simple Mail Transfer Protocol"),
        53: ("DNS", "Domain Name System"),
        80: ("HTTP", "HyperText Transfer Protocol"),
        110: ("POP3", "Post Office Protocol v3"),
        143: ("IMAP", "Internet Message Access Protocol"),
        443: ("HTTPS", "HTTP Secure"),
        445: ("SMB", "Server Message Block"),
        3306: ("MySQL", "MySQL Database"),
        3389: ("RDP", "Remote Desktop Protocol"),
        5432: ("PostgreSQL", "PostgreSQL Database"),
        6379: ("Redis", "Redis Database"),
        8080: ("HTTP Proxy", "HTTP Alternative"),
        27017: ("MongoDB", "MongoDB Database"),
    }
    
    port_info = common_ports.get(input_data.port, ("Unknown", "Port not in common services database"))
    
    return PortAnalyzerOutput(
        port=input_data.port,
        service=port_info[0],
        description=port_info[1]
    )

@api_router.get("/")
async def root():
    return {"message": "VAPA.es API - Sistema de herramientas técnicas"}

# ============= STARTUP EVENT =============

@app.on_event("startup")
async def startup_event():
    # Create admin user if not exists
    admin_email = "admin@vapa.es"
    existing_admin = await db.users.find_one({"email": admin_email})
    
    if not existing_admin:
        admin_user = User(
            email=admin_email,
            password_hash=get_password_hash("admin123"),
            role="admin"
        )
        doc = admin_user.model_dump()
        doc['created_at'] = doc['created_at'].isoformat()
        await db.users.insert_one(doc)
        logging.info(f"Admin user created: {admin_email} / admin123")
    
    # Create initial blog posts if not exist
    posts_count = await db.blog_posts.count_documents({})
    if posts_count == 0:
        initial_posts = [
            {
                "id": str(uuid.uuid4()),
                "title": "Subnetting IPv4: Guía Maestra para Administradores",
                "slug": "guia-subnetting-ipv4-redes",
                "content": "# Subnetting IPv4: Guía Maestra\n\nEl subnetting es una técnica fundamental en la administración de redes...",
                "excerpt": "Aprende a segmentar redes de forma profesional para mejorar la seguridad.",
                "author_id": "system",
                "published": True,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "title": "Seguridad Crítica en Contenedores Docker",
                "slug": "seguridad-docker-servidores-nas",
                "content": "# Seguridad en Docker\n\nLos contenedores Docker requieren consideraciones especiales de seguridad...",
                "excerpt": "Estrategias para proteger despliegues en NAS y OMV.",
                "author_id": "system",
                "published": True,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "title": "JSON vs YAML: ¿Cuál es mejor para tu Proyecto?",
                "slug": "json-vs-yaml-comparativa",
                "content": "# JSON vs YAML\n\nAmbos formatos tienen sus ventajas y desventajas...",
                "excerpt": "Análisis de formatos de serialización de datos.",
                "author_id": "system",
                "published": True,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
        ]
        await db.blog_posts.insert_many(initial_posts)
        logging.info("Initial blog posts created")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)


class VLSMInput(BaseModel):
    network: str
    subnets: List[int]

class VLSMSubnet(BaseModel):
    network: str
    broadcast: str
    first_host: str
    last_host: str
    mask: str
    cidr: int
    hosts: int

class VLSMOutput(BaseModel):
    subnets: List[VLSMSubnet]

class PingInput(BaseModel):
    host: str
    count: int = 4

class PingOutput(BaseModel):
    host: str
    output: str
    success: bool

class PortScanInput(BaseModel):
    host: str
    ports: Optional[str] = "common"

class PortScanResult(BaseModel):
    port: int
    state: str
    service: str

class PortScanOutput(BaseModel):
    host: str
    ip: str
    open_ports: List[PortScanResult]
    scan_time: float

class IPv4toIPv6Input(BaseModel):
    ip: str

class IPv4toIPv6Output(BaseModel):
    ipv4: str
    ipv6_mapped: str
    ipv6_compatible: str
    ipv6_6to4: str

class DNSLookupInput(BaseModel):
    domain: str
    record_type: str = "A"

class DNSLookupOutput(BaseModel):
    domain: str
    record_type: str
    records: List[str]


class VLSMInput(BaseModel):
    network: str
    subnets: List[int]

class VLSMSubnet(BaseModel):
    network: str
    broadcast: str
    first_host: str
    last_host: str
    mask: str
    cidr: int
    hosts: int

class VLSMOutput(BaseModel):
    subnets: List[VLSMSubnet]

class PingInput(BaseModel):
    host: str
    count: int = 4

class PingOutput(BaseModel):
    host: str
    output: str
    success: bool

class PortScanInput(BaseModel):
    host: str
    ports: Optional[str] = "common"

class PortScanResult(BaseModel):
    port: int
    state: str
    service: str

class PortScanOutput(BaseModel):
    host: str
    ip: str
    open_ports: List[PortScanResult]
    scan_time: float

class IPv4toIPv6Input(BaseModel):
    ip: str

class IPv4toIPv6Output(BaseModel):
    ipv4: str
    ipv6_mapped: str
    ipv6_compatible: str
    ipv6_6to4: str

class DNSLookupInput(BaseModel):
    domain: str
    record_type: str = "A"

class DNSLookupOutput(BaseModel):
    domain: str
    record_type: str
    records: List[str]


@api_router.post("/tools/vlsm", response_model=VLSMOutput)
async def vlsm_calculator(input_data: VLSMInput):
    try:
        import ipaddress
        network = ipaddress.ip_network(input_data.network, strict=False)
        subnets_needed = sorted(input_data.subnets, reverse=True)
        result = []
        current_network = network
        for hosts_needed in subnets_needed:
            prefix = 32 - (hosts_needed + 2 - 1).bit_length()
            if prefix < 0:
                prefix = 0
            subnet = ipaddress.ip_network(f"{current_network.network_address}/{prefix}", strict=False)
            if subnet.network_address < current_network.network_address:
                subnet = ipaddress.ip_network(f"{current_network.network_address}/{prefix}", strict=False)
            all_hosts = list(subnet.hosts())
            result.append(VLSMSubnet(
                network=str(subnet.network_address),
                broadcast=str(subnet.broadcast_address),
                first_host=str(all_hosts[0]) if all_hosts else str(subnet.network_address),
                last_host=str(all_hosts[-1]) if all_hosts else str(subnet.broadcast_address),
                mask=str(subnet.netmask),
                cidr=prefix,
                hosts=subnet.num_addresses - 2 if subnet.num_addresses > 2 else 0
            ))
            next_addr = int(subnet.broadcast_address) + 1
            if next_addr > int(network.broadcast_address):
                break
            current_network = ipaddress.ip_network(f"{ipaddress.IPv4Address(next_addr)}/{network.prefixlen}", strict=False)
            current_network = list(ipaddress.ip_network(f"{ipaddress.IPv4Address(next_addr)}/0", strict=False).subnets(new_prefix=network.prefixlen))[0]
            current_network = ipaddress.ip_network(f"{ipaddress.IPv4Address(next_addr)}/{network.prefixlen}", strict=False)
        return VLSMOutput(subnets=result)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/ping", response_model=PingOutput)
async def ping_host(input_data: PingInput):
    import re
    import subprocess
    host = input_data.host.strip()
    if not re.match(r'^[a-zA-Z0-9.\-_]+$', host):
        raise HTTPException(status_code=400, detail="Host inválido")
    count = min(max(input_data.count, 1), 10)
    try:
        result = subprocess.run(
            ["ping", "-c", str(count), "-W", "2", host],
            capture_output=True, text=True, timeout=30
        )
        return PingOutput(
            host=host,
            output=result.stdout + result.stderr,
            success=result.returncode == 0
        )
    except subprocess.TimeoutExpired:
        return PingOutput(host=host, output="Timeout: el host no responde", success=False)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/port-scan", response_model=PortScanOutput)
async def port_scan(input_data: PortScanInput):
    import socket
    import time
    import re
    host = input_data.host.strip()
    if not re.match(r'^[a-zA-Z0-9.\-_]+$', host):
        raise HTTPException(status_code=400, detail="Host inválido")
    common_ports = {
        20: "FTP Data", 21: "FTP", 22: "SSH", 23: "Telnet", 25: "SMTP",
        53: "DNS", 80: "HTTP", 110: "POP3", 143: "IMAP", 443: "HTTPS",
        445: "SMB", 3306: "MySQL", 3389: "RDP", 5432: "PostgreSQL",
        5900: "VNC", 6379: "Redis", 8080: "HTTP-Alt", 8443: "HTTPS-Alt",
        27017: "MongoDB", 9200: "Elasticsearch"
    }
    try:
        ip = socket.gethostbyname(host)
    except socket.gaierror:
        raise HTTPException(status_code=400, detail="No se pudo resolver el host")
    if input_data.ports == "common":
        ports_to_scan = list(common_ports.keys())
    elif "-" in str(input_data.ports):
        start, end = input_data.ports.split("-")
        ports_to_scan = list(range(int(start), min(int(end) + 1, int(start) + 1000)))
    else:
        ports_to_scan = list(common_ports.keys())
    start_time = time.time()
    open_ports = []
    for port in ports_to_scan:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.5)
            result = sock.connect_ex((ip, port))
            if result == 0:
                open_ports.append(PortScanResult(
                    port=port,
                    state="open",
                    service=common_ports.get(port, "unknown")
                ))
            sock.close()
        except:
            pass
    scan_time = round(time.time() - start_time, 2)
    return PortScanOutput(host=host, ip=ip, open_ports=open_ports, scan_time=scan_time)

@api_router.post("/tools/ipv4-to-ipv6", response_model=IPv4toIPv6Output)
async def ipv4_to_ipv6(input_data: IPv4toIPv6Input):
    try:
        import ipaddress
        ipv4 = ipaddress.IPv4Address(input_data.ip)
        octets = str(ipv4).split(".")
        hex_parts = [f"{int(o):02x}" for o in octets]
        ipv6_mapped = f"::ffff:{hex_parts[0]}{hex_parts[1]}:{hex_parts[2]}{hex_parts[3]}"
        ipv6_compatible = f"::{hex_parts[0]}{hex_parts[1]}:{hex_parts[2]}{hex_parts[3]}"
        second_octet = int(octets[1])
        ipv6_6to4 = f"2002:{int(octets[0]):02x}{int(octets[1]):02x}:{int(octets[2]):02x}{int(octets[3]):02x}::1"
        return IPv4toIPv6Output(
            ipv4=str(ipv4),
            ipv6_mapped=ipv6_mapped,
            ipv6_compatible=ipv6_compatible,
            ipv6_6to4=ipv6_6to4
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"IP inválida: {str(e)}")

@api_router.post("/tools/dns-lookup", response_model=DNSLookupOutput)
async def dns_lookup(input_data: DNSLookupInput):
    import subprocess
    import re
    domain = input_data.domain.strip()
    if not re.match(r'^[a-zA-Z0-9.\-_]+$', domain):
        raise HTTPException(status_code=400, detail="Dominio inválido")
    record_type = input_data.record_type.upper()
    if record_type not in ["A", "MX", "CNAME", "TXT", "NS", "AAAA"]:
        raise HTTPException(status_code=400, detail="Tipo de registro no soportado")
    try:
        result = subprocess.run(
            ["dig", "+short", record_type, domain],
            capture_output=True, text=True, timeout=10
        )
        records = [r.strip() for r in result.stdout.strip().split("\n") if r.strip()]
        if not records:
            records = ["No se encontraron registros"]
        return DNSLookupOutput(domain=domain, record_type=record_type, records=records)
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=400, detail="Timeout en la consulta DNS")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

# ============= MODELOS DE PROYECTOS =============

class ProjectMember(BaseModel):
    user_id: str
    role: str = "member"  # owner, member

class Project(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: str
    status: str = "active"  # active, completed, archived
    members: List[ProjectMember] = []
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProjectCreate(BaseModel):
    name: str
    description: str
    status: str = "active"
    member_ids: List[str] = []

class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    member_ids: Optional[List[str]] = None

class ProjectResponse(BaseModel):
    id: str
    name: str
    description: str
    status: str
    members: List[ProjectMember]
    created_by: str
    created_at: datetime
    updated_at: datetime

class Task(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    description: str = ""
    status: str = "pending"  # pending, in_progress, completed
    assigned_to: Optional[str] = None
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TaskCreate(BaseModel):
    title: str
    description: str = ""
    status: str = "pending"
    assigned_to: Optional[str] = None

class TaskUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    assigned_to: Optional[str] = None

class ProjectError(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    description: str = ""
    severity: str = "medium"  # low, medium, high, critical
    status: str = "open"  # open, resolved
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProjectErrorCreate(BaseModel):
    title: str
    description: str = ""
    severity: str = "medium"
    status: str = "open"

class ProjectErrorUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    severity: Optional[str] = None
    status: Optional[str] = None

class Documentation(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    content: str = ""
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class DocumentationCreate(BaseModel):
    title: str
    content: str = ""

class DocumentationUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None

class Report(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    content: str = ""
    report_type: str = "manual"  # manual, auto
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ReportCreate(BaseModel):
    title: str
    content: str = ""
    report_type: str = "manual"

class FileAttachment(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    filename: str
    original_name: str
    file_type: str
    file_size: int
    uploaded_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AddMemberRequest(BaseModel):
    user_id: str
    role: str = "member"  # owner, manager, member

class UpdateMemberRole(BaseModel):
    role: str

@api_router.post("/tools/vlsm", response_model=VLSMOutput)
async def vlsm_calculator(input_data: VLSMInput):
    try:
        import ipaddress
        network = ipaddress.ip_network(input_data.network, strict=False)
        subnets_needed = sorted(input_data.subnets, reverse=True)
        result = []
        current_network = network
        for hosts_needed in subnets_needed:
            prefix = 32 - (hosts_needed + 2 - 1).bit_length()
            if prefix < 0:
                prefix = 0
            subnet = ipaddress.ip_network(f"{current_network.network_address}/{prefix}", strict=False)
            if subnet.network_address < current_network.network_address:
                subnet = ipaddress.ip_network(f"{current_network.network_address}/{prefix}", strict=False)
            all_hosts = list(subnet.hosts())
            result.append(VLSMSubnet(
                network=str(subnet.network_address),
                broadcast=str(subnet.broadcast_address),
                first_host=str(all_hosts[0]) if all_hosts else str(subnet.network_address),
                last_host=str(all_hosts[-1]) if all_hosts else str(subnet.broadcast_address),
                mask=str(subnet.netmask),
                cidr=prefix,
                hosts=subnet.num_addresses - 2 if subnet.num_addresses > 2 else 0
            ))
            next_addr = int(subnet.broadcast_address) + 1
            if next_addr > int(network.broadcast_address):
                break
            current_network = ipaddress.ip_network(f"{ipaddress.IPv4Address(next_addr)}/{network.prefixlen}", strict=False)
            current_network = list(ipaddress.ip_network(f"{ipaddress.IPv4Address(next_addr)}/0", strict=False).subnets(new_prefix=network.prefixlen))[0]
            current_network = ipaddress.ip_network(f"{ipaddress.IPv4Address(next_addr)}/{network.prefixlen}", strict=False)
        return VLSMOutput(subnets=result)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/ping", response_model=PingOutput)
async def ping_host(input_data: PingInput):
    import re
    import subprocess
    host = input_data.host.strip()
    if not re.match(r'^[a-zA-Z0-9.\-_]+$', host):
        raise HTTPException(status_code=400, detail="Host inválido")
    count = min(max(input_data.count, 1), 10)
    try:
        result = subprocess.run(
            ["ping", "-c", str(count), "-W", "2", host],
            capture_output=True, text=True, timeout=30
        )
        return PingOutput(
            host=host,
            output=result.stdout + result.stderr,
            success=result.returncode == 0
        )
    except subprocess.TimeoutExpired:
        return PingOutput(host=host, output="Timeout: el host no responde", success=False)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

@api_router.post("/tools/port-scan", response_model=PortScanOutput)
async def port_scan(input_data: PortScanInput):
    import socket
    import time
    import re
    host = input_data.host.strip()
    if not re.match(r'^[a-zA-Z0-9.\-_]+$', host):
        raise HTTPException(status_code=400, detail="Host inválido")
    common_ports = {
        20: "FTP Data", 21: "FTP", 22: "SSH", 23: "Telnet", 25: "SMTP",
        53: "DNS", 80: "HTTP", 110: "POP3", 143: "IMAP", 443: "HTTPS",
        445: "SMB", 3306: "MySQL", 3389: "RDP", 5432: "PostgreSQL",
        5900: "VNC", 6379: "Redis", 8080: "HTTP-Alt", 8443: "HTTPS-Alt",
        27017: "MongoDB", 9200: "Elasticsearch"
    }
    try:
        ip = socket.gethostbyname(host)
    except socket.gaierror:
        raise HTTPException(status_code=400, detail="No se pudo resolver el host")
    if input_data.ports == "common":
        ports_to_scan = list(common_ports.keys())
    elif "-" in str(input_data.ports):
        start, end = input_data.ports.split("-")
        ports_to_scan = list(range(int(start), min(int(end) + 1, int(start) + 1000)))
    else:
        ports_to_scan = list(common_ports.keys())
    start_time = time.time()
    open_ports = []
    for port in ports_to_scan:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.5)
            result = sock.connect_ex((ip, port))
            if result == 0:
                open_ports.append(PortScanResult(
                    port=port,
                    state="open",
                    service=common_ports.get(port, "unknown")
                ))
            sock.close()
        except:
            pass
    scan_time = round(time.time() - start_time, 2)
    return PortScanOutput(host=host, ip=ip, open_ports=open_ports, scan_time=scan_time)

@api_router.post("/tools/ipv4-to-ipv6", response_model=IPv4toIPv6Output)
async def ipv4_to_ipv6(input_data: IPv4toIPv6Input):
    try:
        import ipaddress
        ipv4 = ipaddress.IPv4Address(input_data.ip)
        octets = str(ipv4).split(".")
        hex_parts = [f"{int(o):02x}" for o in octets]
        ipv6_mapped = f"::ffff:{hex_parts[0]}{hex_parts[1]}:{hex_parts[2]}{hex_parts[3]}"
        ipv6_compatible = f"::{hex_parts[0]}{hex_parts[1]}:{hex_parts[2]}{hex_parts[3]}"
        second_octet = int(octets[1])
        ipv6_6to4 = f"2002:{int(octets[0]):02x}{int(octets[1]):02x}:{int(octets[2]):02x}{int(octets[3]):02x}::1"
        return IPv4toIPv6Output(
            ipv4=str(ipv4),
            ipv6_mapped=ipv6_mapped,
            ipv6_compatible=ipv6_compatible,
            ipv6_6to4=ipv6_6to4
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"IP inválida: {str(e)}")

@api_router.post("/tools/dns-lookup", response_model=DNSLookupOutput)
async def dns_lookup(input_data: DNSLookupInput):
    import subprocess
    import re
    domain = input_data.domain.strip()
    if not re.match(r'^[a-zA-Z0-9.\-_]+$', domain):
        raise HTTPException(status_code=400, detail="Dominio inválido")
    record_type = input_data.record_type.upper()
    if record_type not in ["A", "MX", "CNAME", "TXT", "NS", "AAAA"]:
        raise HTTPException(status_code=400, detail="Tipo de registro no soportado")
    try:
        result = subprocess.run(
            ["dig", "+short", record_type, domain],
            capture_output=True, text=True, timeout=10
        )
        records = [r.strip() for r in result.stdout.strip().split("\n") if r.strip()]
        if not records:
            records = ["No se encontraron registros"]
        return DNSLookupOutput(domain=domain, record_type=record_type, records=records)
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=400, detail="Timeout en la consulta DNS")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

# ============= MODELOS DE PROYECTOS =============

class ProjectMember(BaseModel):
    user_id: str
    role: str = "member"  # owner, member

class Project(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: str
    status: str = "active"  # active, completed, archived
    members: List[ProjectMember] = []
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProjectCreate(BaseModel):
    name: str
    description: str
    status: str = "active"
    member_ids: List[str] = []

class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    member_ids: Optional[List[str]] = None

class ProjectResponse(BaseModel):
    id: str
    name: str
    description: str
    status: str
    members: List[ProjectMember]
    created_by: str
    created_at: datetime
    updated_at: datetime

class Task(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    description: str = ""
    status: str = "pending"  # pending, in_progress, completed
    assigned_to: Optional[str] = None
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TaskCreate(BaseModel):
    title: str
    description: str = ""
    status: str = "pending"
    assigned_to: Optional[str] = None

class TaskUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    assigned_to: Optional[str] = None

class ProjectError(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    description: str = ""
    severity: str = "medium"  # low, medium, high, critical
    status: str = "open"  # open, resolved
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProjectErrorCreate(BaseModel):
    title: str
    description: str = ""
    severity: str = "medium"
    status: str = "open"

class ProjectErrorUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    severity: Optional[str] = None
    status: Optional[str] = None

class Documentation(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    content: str = ""
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class DocumentationCreate(BaseModel):
    title: str
    content: str = ""

class DocumentationUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None

class Report(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    title: str
    content: str = ""
    report_type: str = "manual"  # manual, auto
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ReportCreate(BaseModel):
    title: str
    content: str = ""
    report_type: str = "manual"

class FileAttachment(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    filename: str
    original_name: str
    file_type: str
    file_size: int
    uploaded_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AddMemberRequest(BaseModel):
    user_id: str
    role: str = "member"  # owner, manager, member

class UpdateMemberRole(BaseModel):
    role: str

@api_router.get("/projects", response_model=List[ProjectResponse])
async def get_projects(current_user: User = Depends(get_current_user)):
    if current_user.role == "admin":
        projects = await db.projects.find({}, {"_id": 0}).to_list(100)
    else:
        projects = await db.projects.find(
            {"$or": [
                {"created_by": current_user.id},
                {"members": {"$elemMatch": {"user_id": current_user.id}}}
            ]}, {"_id": 0}
        ).to_list(100)
    result = []
    for p in projects:
        for field in ["created_at", "updated_at"]:
            if isinstance(p.get(field), str):
                p[field] = datetime.fromisoformat(p[field])
        result.append(ProjectResponse(**p))
    return result

@api_router.post("/projects", response_model=ProjectResponse)
async def create_project(data: ProjectCreate, current_user: User = Depends(get_current_active_editor)):
    members = [ProjectMember(user_id=uid) for uid in data.member_ids]
    project = Project(
        name=data.name,
        description=data.description,
        status=data.status,
        members=members,
        created_by=current_user.id
    )
    doc = project.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    doc["members"] = [m.model_dump() for m in members]
    await db.projects.insert_one(doc)
    return ProjectResponse(**project.model_dump())

@api_router.get("/projects/{project_id}", response_model=ProjectResponse)
async def get_project(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    for field in ["created_at", "updated_at"]:
        if isinstance(project.get(field), str):
            project[field] = datetime.fromisoformat(project[field])
    return ProjectResponse(**project)

@api_router.put("/projects/{project_id}", response_model=ProjectResponse)
async def update_project(project_id: str, data: ProjectUpdate, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    if project["created_by"] != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes permisos para editar este proyecto")
    update = data.model_dump(exclude_unset=True)
    if "member_ids" in update:
        update["members"] = [{"user_id": uid, "role": "member"} for uid in update.pop("member_ids")]
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.projects.update_one({"id": project_id}, {"$set": update})
    updated = await db.projects.find_one({"id": project_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return ProjectResponse(**updated)

@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str, current_user: User = Depends(get_current_active_admin)):
    result = await db.projects.delete_one({"id": project_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    await db.tasks.delete_many({"project_id": project_id})
    await db.project_errors.delete_many({"project_id": project_id})
    await db.documentation.delete_many({"project_id": project_id})
    await db.reports.delete_many({"project_id": project_id})
    await db.files.delete_many({"project_id": project_id})
    return {"message": "Proyecto eliminado"}

# ============= RUTAS DE TAREAS =============

@api_router.get("/projects/{project_id}/tasks", response_model=List[Task])
async def get_tasks(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    tasks = await db.tasks.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for t in tasks:
        for field in ["created_at", "updated_at"]:
            if isinstance(t.get(field), str):
                t[field] = datetime.fromisoformat(t[field])
    return tasks

@api_router.post("/projects/{project_id}/tasks", response_model=Task)
async def create_task(project_id: str, data: TaskCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    task = Task(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = task.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.tasks.insert_one(doc)
    return task

@api_router.put("/projects/{project_id}/tasks/{task_id}", response_model=Task)
async def update_task(project_id: str, task_id: str, data: TaskUpdate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    update = data.model_dump(exclude_unset=True)
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.tasks.update_one({"id": task_id, "project_id": project_id}, {"$set": update})
    updated = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return Task(**updated)

@api_router.delete("/projects/{project_id}/tasks/{task_id}")
async def delete_task(project_id: str, task_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.tasks.delete_one({"id": task_id, "project_id": project_id})
    return {"message": "Tarea eliminada"}

# ============= RUTAS DE ERRORES =============

@api_router.get("/projects/{project_id}/errors", response_model=List[ProjectError])
async def get_errors(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    errors = await db.project_errors.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for e in errors:
        for field in ["created_at", "updated_at"]:
            if isinstance(e.get(field), str):
                e[field] = datetime.fromisoformat(e[field])
    return errors

@api_router.post("/projects/{project_id}/errors", response_model=ProjectError)
async def create_error(project_id: str, data: ProjectErrorCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    error = ProjectError(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = error.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.project_errors.insert_one(doc)
    return error

@api_router.put("/projects/{project_id}/errors/{error_id}", response_model=ProjectError)
async def update_error(project_id: str, error_id: str, data: ProjectErrorUpdate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    update = data.model_dump(exclude_unset=True)
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.project_errors.update_one({"id": error_id, "project_id": project_id}, {"$set": update})
    updated = await db.project_errors.find_one({"id": error_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return ProjectError(**updated)

@api_router.delete("/projects/{project_id}/errors/{error_id}")
async def delete_error(project_id: str, error_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.project_errors.delete_one({"id": error_id, "project_id": project_id})
    return {"message": "Error eliminado"}

# ============= RUTAS DE DOCUMENTACION =============

@api_router.get("/projects/{project_id}/docs", response_model=List[Documentation])
async def get_docs(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    docs = await db.documentation.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for d in docs:
        for field in ["created_at", "updated_at"]:
            if isinstance(d.get(field), str):
                d[field] = datetime.fromisoformat(d[field])
    return docs

@api_router.post("/projects/{project_id}/docs", response_model=Documentation)
async def create_doc(project_id: str, data: DocumentationCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    doc_obj = Documentation(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = doc_obj.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.documentation.insert_one(doc)
    return doc_obj

@api_router.put("/projects/{project_id}/docs/{doc_id}", response_model=Documentation)
async def update_doc(project_id: str, doc_id: str, data: DocumentationUpdate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    update = data.model_dump(exclude_unset=True)
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.documentation.update_one({"id": doc_id, "project_id": project_id}, {"$set": update})
    updated = await db.documentation.find_one({"id": doc_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return Documentation(**updated)

@api_router.delete("/projects/{project_id}/docs/{doc_id}")
async def delete_doc(project_id: str, doc_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.documentation.delete_one({"id": doc_id, "project_id": project_id})
    return {"message": "Documento eliminado"}

# ============= RUTAS DE INFORMES =============

@api_router.get("/projects/{project_id}/reports", response_model=List[Report])
async def get_reports(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    reports = await db.reports.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for r in reports:
        if isinstance(r.get("created_at"), str):
            r["created_at"] = datetime.fromisoformat(r["created_at"])
    return reports

@api_router.post("/projects/{project_id}/reports", response_model=Report)
async def create_report(project_id: str, data: ReportCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    if data.report_type == "auto":
        project = await db.projects.find_one({"id": project_id}, {"_id": 0})
        tasks = await db.tasks.find({"project_id": project_id}, {"_id": 0}).to_list(100)
        errors = await db.project_errors.find({"project_id": project_id}, {"_id": 0}).to_list(100)
        docs = await db.documentation.find({"project_id": project_id}, {"_id": 0}).to_list(100)
        tasks_pending = len([t for t in tasks if t["status"] == "pending"])
        tasks_progress = len([t for t in tasks if t["status"] == "in_progress"])
        tasks_done = len([t for t in tasks if t["status"] == "completed"])
        errors_open = len([e for e in errors if e["status"] == "open"])
        errors_resolved = len([e for e in errors if e["status"] == "resolved"])
        content = f"""# Informe Automático: {project["name"]}

## Descripción
{project["description"]}

## Estado del Proyecto
Estado actual: {project["status"]}

## Resumen de Tareas
- Pendientes: {tasks_pending}
- En progreso: {tasks_progress}
- Completadas: {tasks_done}
- Total: {len(tasks)}

## Resumen de Errores
- Abiertos: {errors_open}
- Resueltos: {errors_resolved}
- Total: {len(errors)}

## Documentación
- Documentos disponibles: {len(docs)}

## Fecha de generación
{datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M")} UTC
"""
        data.content = content
    report = Report(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = report.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.reports.insert_one(doc)
    return report

@api_router.delete("/projects/{project_id}/reports/{report_id}")
async def delete_report(project_id: str, report_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.reports.delete_one({"id": report_id, "project_id": project_id})
    return {"message": "Informe eliminado"}

@api_router.get("/projects/{project_id}/reports/{report_id}/download/{format}")
async def download_report(project_id: str, report_id: str, format: str, current_user: User = Depends(get_current_user)):
    from fastapi.responses import StreamingResponse
    import io
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    report = await db.reports.find_one({"id": report_id, "project_id": project_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Informe no encontrado")
    if format == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        title = document.add_heading(report["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x20, 0x5C, 0x00)
        document.add_paragraph(f'Proyecto: {project_id}')
        document.add_paragraph(f'Fecha: {report["created_at"]}')
        document.add_paragraph("")
        for line in report["content"].split("\n"):
            if line.startswith("# "):
                p = document.add_heading(line[2:], 1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x20, 0x5C, 0x00)
            elif line.startswith("## "):
                p = document.add_heading(line[3:], 2)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x4A, 0x8C, 0x00)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                document.add_paragraph(line)
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        filename = f"{report['title'].replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    elif format == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        lime = HexColor("#cdff00")
        dark = HexColor("#1a1a1a")
        styles.add(ParagraphStyle(name="VapaTitle", fontSize=20, textColor=HexColor("#cdff00"), spaceAfter=12, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name="VapaH1", fontSize=14, textColor=HexColor("#cdff00"), spaceAfter=8, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name="VapaH2", fontSize=12, textColor=HexColor("#88cc00"), spaceAfter=6, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name="VapaBody", fontSize=10, textColor=HexColor("#333333"), spaceAfter=4))
        story = []
        story.append(Paragraph(report["title"], styles["VapaTitle"]))
        story.append(Spacer(1, 0.5*cm))
        for line in report["content"].split("\n"):
            if line.startswith("# "):
                story.append(Paragraph(line[2:], styles["VapaH1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["VapaH2"]))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {line[2:]}", styles["VapaBody"]))
            elif line.strip():
                story.append(Paragraph(line, styles["VapaBody"]))
            else:
                story.append(Spacer(1, 0.3*cm))
        doc.build(story)
        buffer.seek(0)
        filename = f"{report['title'].replace(' ', '_')}.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    else:
        raise HTTPException(status_code=400, detail="Formato no soportado. Usa 'pdf' o 'docx'")

# ============= RUTAS DE MIEMBROS =============

@api_router.get("/projects/{project_id}/members")
async def get_members(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    members = project.get("members", [])
    result = []
    for m in members:
        user = await db.users.find_one({"id": m["user_id"]}, {"_id": 0})
        if user:
            result.append({"user_id": m["user_id"], "email": user["email"], "role": m.get("role", "member")})
    return result

@api_router.post("/projects/{project_id}/members")
async def add_member(project_id: str, data: AddMemberRequest, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    is_owner = project["created_by"] == current_user.id
    is_manager = any(m["user_id"] == current_user.id and m.get("role") in ["owner", "manager"] for m in project.get("members", []))
    if not is_owner and not is_manager and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes permisos para añadir miembros")
    existing = any(m["user_id"] == data.user_id for m in project.get("members", []))
    if existing:
        raise HTTPException(status_code=400, detail="El usuario ya es miembro del proyecto")
    await db.projects.update_one(
        {"id": project_id},
        {"$push": {"members": {"user_id": data.user_id, "role": data.role}}}
    )
    return {"message": "Miembro añadido"}

@api_router.put("/projects/{project_id}/members/{user_id}")
async def update_member_role(project_id: str, user_id: str, data: UpdateMemberRole, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    is_owner = project["created_by"] == current_user.id
    if not is_owner and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Solo el propietario o admin puede cambiar roles")
    await db.projects.update_one(
        {"id": project_id, "members.user_id": user_id},
        {"$set": {"members.$.role": data.role}}
    )
    return {"message": "Rol actualizado"}

@api_router.delete("/projects/{project_id}/members/{user_id}")
async def remove_member(project_id: str, user_id: str, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    is_owner = project["created_by"] == current_user.id
    is_manager = any(m["user_id"] == current_user.id and m.get("role") in ["owner", "manager"] for m in project.get("members", []))
    if not is_owner and not is_manager and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes permisos para eliminar miembros")
    await db.projects.update_one(
        {"id": project_id},
        {"$pull": {"members": {"user_id": user_id}}}
    )
    return {"message": "Miembro eliminado"}

@api_router.get("/projects/{project_id}/available-users")
async def get_available_users(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    member_ids = [m["user_id"] for m in project.get("members", [])]
    member_ids.append(project["created_by"])
    all_users = await db.users.find({}, {"_id": 0}).to_list(100)
    available = [{"id": u["id"], "email": u["email"]} for u in all_users if u["id"] not in member_ids]
    return available


# ============= RUTAS DE ARCHIVOS =============

import shutil

UPLOAD_DIR = "/app/uploads"

@api_router.post("/projects/{project_id}/files")
async def upload_file(project_id: str, file: UploadFile, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    os.makedirs(f"{UPLOAD_DIR}/{project_id}", exist_ok=True)
    file_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename)[1]
    saved_name = f"{file_id}{ext}"
    file_path = f"{UPLOAD_DIR}/{project_id}/{saved_name}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    file_size = os.path.getsize(file_path)
    attachment = FileAttachment(
        project_id=project_id,
        filename=saved_name,
        original_name=file.filename,
        file_type=file.content_type,
        file_size=file_size,
        uploaded_by=current_user.id,
        category=category
    )
    doc = attachment.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.files.insert_one(doc)
    return {"id": attachment.id, "filename": file.filename, "size": file_size}

@api_router.get("/projects/{project_id}/files")
async def get_files(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    files = await db.files.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for f in files:
        if isinstance(f.get("created_at"), str):
            f["created_at"] = datetime.fromisoformat(f["created_at"])
    return files

@api_router.get("/projects/{project_id}/files/{file_id}/download")
async def download_file(project_id: str, file_id: str, current_user: User = Depends(get_current_user)):
    from fastapi.responses import FileResponse
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    file_doc = await db.files.find_one({"id": file_id, "project_id": project_id}, {"_id": 0})
    if not file_doc:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    file_path = f"{UPLOAD_DIR}/{project_id}/{file_doc['filename']}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado en disco")
    return FileResponse(file_path, filename=file_doc["original_name"])

@api_router.delete("/projects/{project_id}/files/{file_id}")
async def delete_file(project_id: str, file_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    file_doc = await db.files.find_one({"id": file_id, "project_id": project_id}, {"_id": 0})
    if file_doc:
        file_path = f"{UPLOAD_DIR}/{project_id}/{file_doc['filename']}"
        if os.path.exists(file_path):
            os.remove(file_path)
    await db.files.delete_one({"id": file_id, "project_id": project_id})
    return {"message": "Archivo eliminado"}

# Include router
app.include_router(api_router)

# Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

class AddMemberRequest(BaseModel):
    user_id: str
    role: str = "member"  # owner, manager, member

class UpdateMemberRole(BaseModel):
    role: str

@api_router.get("/projects", response_model=List[ProjectResponse])
async def get_projects(current_user: User = Depends(get_current_user)):
    if current_user.role == "admin":
        projects = await db.projects.find({}, {"_id": 0}).to_list(100)
    else:
        projects = await db.projects.find(
            {"$or": [
                {"created_by": current_user.id},
                {"members": {"$elemMatch": {"user_id": current_user.id}}}
            ]}, {"_id": 0}
        ).to_list(100)
    result = []
    for p in projects:
        for field in ["created_at", "updated_at"]:
            if isinstance(p.get(field), str):
                p[field] = datetime.fromisoformat(p[field])
        result.append(ProjectResponse(**p))
    return result

@api_router.post("/projects", response_model=ProjectResponse)
async def create_project(data: ProjectCreate, current_user: User = Depends(get_current_active_editor)):
    members = [ProjectMember(user_id=uid) for uid in data.member_ids]
    project = Project(
        name=data.name,
        description=data.description,
        status=data.status,
        members=members,
        created_by=current_user.id
    )
    doc = project.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    doc["members"] = [m.model_dump() for m in members]
    await db.projects.insert_one(doc)
    return ProjectResponse(**project.model_dump())

@api_router.get("/projects/{project_id}", response_model=ProjectResponse)
async def get_project(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    for field in ["created_at", "updated_at"]:
        if isinstance(project.get(field), str):
            project[field] = datetime.fromisoformat(project[field])
    return ProjectResponse(**project)

@api_router.put("/projects/{project_id}", response_model=ProjectResponse)
async def update_project(project_id: str, data: ProjectUpdate, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    if project["created_by"] != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes permisos para editar este proyecto")
    update = data.model_dump(exclude_unset=True)
    if "member_ids" in update:
        update["members"] = [{"user_id": uid, "role": "member"} for uid in update.pop("member_ids")]
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.projects.update_one({"id": project_id}, {"$set": update})
    updated = await db.projects.find_one({"id": project_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return ProjectResponse(**updated)

@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str, current_user: User = Depends(get_current_active_admin)):
    result = await db.projects.delete_one({"id": project_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    await db.tasks.delete_many({"project_id": project_id})
    await db.project_errors.delete_many({"project_id": project_id})
    await db.documentation.delete_many({"project_id": project_id})
    await db.reports.delete_many({"project_id": project_id})
    await db.files.delete_many({"project_id": project_id})
    return {"message": "Proyecto eliminado"}

# ============= RUTAS DE TAREAS =============

@api_router.get("/projects/{project_id}/tasks", response_model=List[Task])
async def get_tasks(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    tasks = await db.tasks.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for t in tasks:
        for field in ["created_at", "updated_at"]:
            if isinstance(t.get(field), str):
                t[field] = datetime.fromisoformat(t[field])
    return tasks

@api_router.post("/projects/{project_id}/tasks", response_model=Task)
async def create_task(project_id: str, data: TaskCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    task = Task(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = task.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.tasks.insert_one(doc)
    return task

@api_router.put("/projects/{project_id}/tasks/{task_id}", response_model=Task)
async def update_task(project_id: str, task_id: str, data: TaskUpdate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    update = data.model_dump(exclude_unset=True)
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.tasks.update_one({"id": task_id, "project_id": project_id}, {"$set": update})
    updated = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return Task(**updated)

@api_router.delete("/projects/{project_id}/tasks/{task_id}")
async def delete_task(project_id: str, task_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.tasks.delete_one({"id": task_id, "project_id": project_id})
    return {"message": "Tarea eliminada"}

# ============= RUTAS DE ERRORES =============

@api_router.get("/projects/{project_id}/errors", response_model=List[ProjectError])
async def get_errors(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    errors = await db.project_errors.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for e in errors:
        for field in ["created_at", "updated_at"]:
            if isinstance(e.get(field), str):
                e[field] = datetime.fromisoformat(e[field])
    return errors

@api_router.post("/projects/{project_id}/errors", response_model=ProjectError)
async def create_error(project_id: str, data: ProjectErrorCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    error = ProjectError(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = error.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.project_errors.insert_one(doc)
    return error

@api_router.put("/projects/{project_id}/errors/{error_id}", response_model=ProjectError)
async def update_error(project_id: str, error_id: str, data: ProjectErrorUpdate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    update = data.model_dump(exclude_unset=True)
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.project_errors.update_one({"id": error_id, "project_id": project_id}, {"$set": update})
    updated = await db.project_errors.find_one({"id": error_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return ProjectError(**updated)

@api_router.delete("/projects/{project_id}/errors/{error_id}")
async def delete_error(project_id: str, error_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.project_errors.delete_one({"id": error_id, "project_id": project_id})
    return {"message": "Error eliminado"}

# ============= RUTAS DE DOCUMENTACION =============

@api_router.get("/projects/{project_id}/docs", response_model=List[Documentation])
async def get_docs(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    docs = await db.documentation.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for d in docs:
        for field in ["created_at", "updated_at"]:
            if isinstance(d.get(field), str):
                d[field] = datetime.fromisoformat(d[field])
    return docs

@api_router.post("/projects/{project_id}/docs", response_model=Documentation)
async def create_doc(project_id: str, data: DocumentationCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    doc_obj = Documentation(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = doc_obj.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.documentation.insert_one(doc)
    return doc_obj

@api_router.put("/projects/{project_id}/docs/{doc_id}", response_model=Documentation)
async def update_doc(project_id: str, doc_id: str, data: DocumentationUpdate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    update = data.model_dump(exclude_unset=True)
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.documentation.update_one({"id": doc_id, "project_id": project_id}, {"$set": update})
    updated = await db.documentation.find_one({"id": doc_id}, {"_id": 0})
    for field in ["created_at", "updated_at"]:
        if isinstance(updated.get(field), str):
            updated[field] = datetime.fromisoformat(updated[field])
    return Documentation(**updated)

@api_router.delete("/projects/{project_id}/docs/{doc_id}")
async def delete_doc(project_id: str, doc_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.documentation.delete_one({"id": doc_id, "project_id": project_id})
    return {"message": "Documento eliminado"}

# ============= RUTAS DE INFORMES =============

@api_router.get("/projects/{project_id}/reports", response_model=List[Report])
async def get_reports(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    reports = await db.reports.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for r in reports:
        if isinstance(r.get("created_at"), str):
            r["created_at"] = datetime.fromisoformat(r["created_at"])
    return reports

@api_router.post("/projects/{project_id}/reports", response_model=Report)
async def create_report(project_id: str, data: ReportCreate, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    if data.report_type == "auto":
        project = await db.projects.find_one({"id": project_id}, {"_id": 0})
        tasks = await db.tasks.find({"project_id": project_id}, {"_id": 0}).to_list(100)
        errors = await db.project_errors.find({"project_id": project_id}, {"_id": 0}).to_list(100)
        docs = await db.documentation.find({"project_id": project_id}, {"_id": 0}).to_list(100)
        tasks_pending = len([t for t in tasks if t["status"] == "pending"])
        tasks_progress = len([t for t in tasks if t["status"] == "in_progress"])
        tasks_done = len([t for t in tasks if t["status"] == "completed"])
        errors_open = len([e for e in errors if e["status"] == "open"])
        errors_resolved = len([e for e in errors if e["status"] == "resolved"])
        content = f"""# Informe Automático: {project["name"]}

## Descripción
{project["description"]}

## Estado del Proyecto
Estado actual: {project["status"]}

## Resumen de Tareas
- Pendientes: {tasks_pending}
- En progreso: {tasks_progress}
- Completadas: {tasks_done}
- Total: {len(tasks)}

## Resumen de Errores
- Abiertos: {errors_open}
- Resueltos: {errors_resolved}
- Total: {len(errors)}

## Documentación
- Documentos disponibles: {len(docs)}

## Fecha de generación
{datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M")} UTC
"""
        data.content = content
    report = Report(project_id=project_id, created_by=current_user.id, **data.model_dump())
    doc = report.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.reports.insert_one(doc)
    return report

@api_router.delete("/projects/{project_id}/reports/{report_id}")
async def delete_report(project_id: str, report_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    await db.reports.delete_one({"id": report_id, "project_id": project_id})
    return {"message": "Informe eliminado"}

@api_router.get("/projects/{project_id}/reports/{report_id}/download/{format}")
async def download_report(project_id: str, report_id: str, format: str, current_user: User = Depends(get_current_user)):
    from fastapi.responses import StreamingResponse
    import io
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    report = await db.reports.find_one({"id": report_id, "project_id": project_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Informe no encontrado")
    if format == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        title = document.add_heading(report["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x20, 0x5C, 0x00)
        document.add_paragraph(f'Proyecto: {project_id}')
        document.add_paragraph(f'Fecha: {report["created_at"]}')
        document.add_paragraph("")
        for line in report["content"].split("\n"):
            if line.startswith("# "):
                p = document.add_heading(line[2:], 1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x20, 0x5C, 0x00)
            elif line.startswith("## "):
                p = document.add_heading(line[3:], 2)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x4A, 0x8C, 0x00)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                document.add_paragraph(line)
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        filename = f"{report['title'].replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    elif format == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        lime = HexColor("#cdff00")
        dark = HexColor("#1a1a1a")
        styles.add(ParagraphStyle(name="VapaTitle", fontSize=20, textColor=HexColor("#cdff00"), spaceAfter=12, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name="VapaH1", fontSize=14, textColor=HexColor("#cdff00"), spaceAfter=8, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name="VapaH2", fontSize=12, textColor=HexColor("#88cc00"), spaceAfter=6, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name="VapaBody", fontSize=10, textColor=HexColor("#333333"), spaceAfter=4))
        story = []
        story.append(Paragraph(report["title"], styles["VapaTitle"]))
        story.append(Spacer(1, 0.5*cm))
        for line in report["content"].split("\n"):
            if line.startswith("# "):
                story.append(Paragraph(line[2:], styles["VapaH1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["VapaH2"]))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {line[2:]}", styles["VapaBody"]))
            elif line.strip():
                story.append(Paragraph(line, styles["VapaBody"]))
            else:
                story.append(Spacer(1, 0.3*cm))
        doc.build(story)
        buffer.seek(0)
        filename = f"{report['title'].replace(' ', '_')}.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    else:
        raise HTTPException(status_code=400, detail="Formato no soportado. Usa 'pdf' o 'docx'")

# ============= RUTAS DE MIEMBROS =============

@api_router.get("/projects/{project_id}/members")
async def get_members(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    members = project.get("members", [])
    result = []
    for m in members:
        user = await db.users.find_one({"id": m["user_id"]}, {"_id": 0})
        if user:
            result.append({"user_id": m["user_id"], "email": user["email"], "role": m.get("role", "member")})
    return result

@api_router.post("/projects/{project_id}/members")
async def add_member(project_id: str, data: AddMemberRequest, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    is_owner = project["created_by"] == current_user.id
    is_manager = any(m["user_id"] == current_user.id and m.get("role") in ["owner", "manager"] for m in project.get("members", []))
    if not is_owner and not is_manager and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes permisos para añadir miembros")
    existing = any(m["user_id"] == data.user_id for m in project.get("members", []))
    if existing:
        raise HTTPException(status_code=400, detail="El usuario ya es miembro del proyecto")
    await db.projects.update_one(
        {"id": project_id},
        {"$push": {"members": {"user_id": data.user_id, "role": data.role}}}
    )
    return {"message": "Miembro añadido"}

@api_router.put("/projects/{project_id}/members/{user_id}")
async def update_member_role(project_id: str, user_id: str, data: UpdateMemberRole, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    is_owner = project["created_by"] == current_user.id
    if not is_owner and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Solo el propietario o admin puede cambiar roles")
    await db.projects.update_one(
        {"id": project_id, "members.user_id": user_id},
        {"$set": {"members.$.role": data.role}}
    )
    return {"message": "Rol actualizado"}

@api_router.delete("/projects/{project_id}/members/{user_id}")
async def remove_member(project_id: str, user_id: str, current_user: User = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    is_owner = project["created_by"] == current_user.id
    is_manager = any(m["user_id"] == current_user.id and m.get("role") in ["owner", "manager"] for m in project.get("members", []))
    if not is_owner and not is_manager and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes permisos para eliminar miembros")
    await db.projects.update_one(
        {"id": project_id},
        {"$pull": {"members": {"user_id": user_id}}}
    )
    return {"message": "Miembro eliminado"}

@api_router.get("/projects/{project_id}/available-users")
async def get_available_users(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    member_ids = [m["user_id"] for m in project.get("members", [])]
    member_ids.append(project["created_by"])
    all_users = await db.users.find({}, {"_id": 0}).to_list(100)
    available = [{"id": u["id"], "email": u["email"]} for u in all_users if u["id"] not in member_ids]
    return available


# ============= RUTAS DE ARCHIVOS =============

import shutil

UPLOAD_DIR = "/app/uploads"

@api_router.post("/projects/{project_id}/files")
async def upload_file(project_id: str, file: UploadFile, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    os.makedirs(f"{UPLOAD_DIR}/{project_id}", exist_ok=True)
    file_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename)[1]
    saved_name = f"{file_id}{ext}"
    file_path = f"{UPLOAD_DIR}/{project_id}/{saved_name}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    file_size = os.path.getsize(file_path)
    attachment = FileAttachment(
        project_id=project_id,
        filename=saved_name,
        original_name=file.filename,
        file_type=file.content_type,
        file_size=file_size,
        uploaded_by=current_user.id,
        category=category
    )
    doc = attachment.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.files.insert_one(doc)
    return {"id": attachment.id, "filename": file.filename, "size": file_size}

@api_router.get("/projects/{project_id}/files")
async def get_files(project_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    files = await db.files.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    for f in files:
        if isinstance(f.get("created_at"), str):
            f["created_at"] = datetime.fromisoformat(f["created_at"])
    return files

@api_router.get("/projects/{project_id}/files/{file_id}/download")
async def download_file(project_id: str, file_id: str, current_user: User = Depends(get_current_user)):
    from fastapi.responses import FileResponse
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    file_doc = await db.files.find_one({"id": file_id, "project_id": project_id}, {"_id": 0})
    if not file_doc:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    file_path = f"{UPLOAD_DIR}/{project_id}/{file_doc['filename']}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado en disco")
    return FileResponse(file_path, filename=file_doc["original_name"])

@api_router.delete("/projects/{project_id}/files/{file_id}")
async def delete_file(project_id: str, file_id: str, current_user: User = Depends(get_current_user)):
    if not await check_project_member(project_id, current_user.id) and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
    file_doc = await db.files.find_one({"id": file_id, "project_id": project_id}, {"_id": 0})
    if file_doc:
        file_path = f"{UPLOAD_DIR}/{project_id}/{file_doc['filename']}"
        if os.path.exists(file_path):
            os.remove(file_path)
    await db.files.delete_one({"id": file_id, "project_id": project_id})
    return {"message": "Archivo eliminado"}

# Include router
app.include_router(api_router)

# Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

